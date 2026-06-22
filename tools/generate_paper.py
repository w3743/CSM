"""生成 CSM 学术论文 Word 文档。"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── 样式设置 ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 标题样式
for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = 'Times New Roman'
    h_font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h_font.size = Pt(14)
        h_font.bold = True
    elif level == 2:
        h_font.size = Pt(12)
        h_font.bold = True
    else:
        h_font.size = Pt(11)
        h_font.bold = True

# ── 辅助函数 ──────────────────────────────────────────────
def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p

# ════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('连续强度记忆系统 (CSM)\n— 基于间隔效应与自适应进化的\nLLM Agent 长期记忆模型')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()
sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_para.add_run('Continuous Strength Memory for LLM Agents:\nA Spacing-Effect-Driven Memory Model with Adaptive Evolution')
run.font.size = Pt(13)
run.font.italic = True
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_para.add_run('CSM Agent Project\n').font.size = Pt(12)
info_para.add_run('https://github.com/w3743/CSM\n').font.size = Pt(11)
info_para.add_run('2026 年 6 月').font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 摘要
# ════════════════════════════════════════════════════════════
doc.add_heading('摘  要', level=1)

add_body(
    '大型语言模型（LLM）驱动的 AI Agent 面临跨会话状态遗忘的固有问题。'
    '现有方案或依赖全量上下文窗口（成本随会话数线性增长），'
    '或采用固定规则的记忆系统（缺乏自适应能力）。'
    '本文提出连续强度记忆系统（Continuous Strength Memory, CSM），'
    '一种为 LLM Agent 设计的轻量级长期记忆子系统，核心理念为"从两条简单规则中涌现智能行为"。'
)

add_body(
    'CSM 的记忆模型基于两个基本力：指数衰减（模拟遗忘）和间隔强化（模拟复习）。'
    '系统借鉴墨墨背单词/FSRS（Free Spaced Repetition Scheduler）的间隔效应理论，'
    '提出了时间感知的强化函数：强化增益与当前可回忆概率 R 呈非线性关系，'
    '最大增益出现在 R ≈ 0.5 的半遗忘状态。每条记忆拥有独立的衰减率（decay_rate）、'
    '检索偏置（boost）和信任度（trust），由自适应进化引擎根据 LLM 的使用/无视/纠正反馈自动调整。'
    '检索采用语义嵌入（BGE-large-zh-v1.5）与 FTS5 关键词的混合策略，'
    '最终得分为语义相似度、当前强度与经验偏置的乘积。'
    '系统以零外部依赖的 SQLite + HTTP Sidecar 架构实现，'
    '提供 REST API 和 Web 管理控制台，可通过 pi agent 扩展自动集成。'
    '实验表明，CSM 的间隔强化函数相比固定增益强化在长期记忆保留率上具有显著优势。'
)

# 关键词
p = doc.add_paragraph()
run = p.add_run('关键词：')
run.font.bold = True
p.add_run('连续强度记忆；间隔效应；间隔重复；LLM Agent；自适应进化；语义检索；SQLite')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. 引言
# ════════════════════════════════════════════════════════════
doc.add_heading('1. 引言', level=1)

add_body(
    '随着大型语言模型（LLM）在编码助手、对话系统等领域的广泛应用，'
    '跨会话持久化记忆成为限制 Agent 长期可用性的核心瓶颈。'
    '典型的 LLM Agent（如 Cursor、Claude Code、pi coding agent）将每轮对话的上下文窗口'
    '全部传输给模型，导致以下问题：（1）上下文窗口随会话增长而线性膨胀，'
    'Token 成本不可持续；（2）Agent 无法记住跨会话的用户偏好、项目约定和纠正历史；'
    '（3）长上下文中的信息检索精度随距离衰减 [Liu et al., 2023]。'
)

add_body(
    '间隔重复系统（Spaced Repetition System, SRS）在教育领域取得了巨大成功。'
    '从 Ebbinghaus (1885) 的遗忘曲线到 SuperMemo 的 SM-2 算法 [Wozniak, 1990]，'
    '再到墨墨背单词的 FSRS 模型 [Ye, 2022]，间隔重复理论已经证明：'
    '适时的复习可以指数级延长记忆保持时间。'
    '然而，这些系统均为人类学习者设计，依赖显式的"评分"反馈（Again/Hard/Good/Easy），'
    '无法直接应用于 LLM Agent 的隐式交互场景。'
)

add_body(
    '本文的贡献包括：（1）将间隔重复理论适配到 LLM Agent 记忆场景，'
    '提出了基于隐式反馈（used/ignored/corrected）的自适应记忆模型；'
    '（2）设计了时间感知的强化函数，使复习时机（而不仅仅是复习次数）影响长期记忆强度；'
    '（3）构建了完整的记忆生命周期——检索、强化、进化、归档——并以零外部依赖实现；'
    '（4）开发了可 pip 安装的开源系统和 pi Agent 自动集成扩展。'
)

# ════════════════════════════════════════════════════════════
# 2. 核心数学模型
# ════════════════════════════════════════════════════════════
doc.add_heading('2. 核心数学模型', level=1)

doc.add_heading('2.1 遗忘函数', level=2)

add_body(
    'CSM 采用指数衰减模型描述记忆强度的自然退化。'
    '记一条记忆在 t 天前的存储强度为 s₀，当前可回忆概率（Retrievability）为 R(t)，'
    '则有：'
)

add_equation('R(t) = s₀ · e^(-d · t)')

add_body(
    '其中 d 为衰减率（decay_rate），表示强度每天衰减的速度。'
    '该公式与 Ebbinghaus 遗忘曲线及 FSRS 的 R(t) = 2^(-t/S) 在数学上等价'
    '（令 d = ln 2 / S，其中 S 为半衰期）。'
    '与 FSRS 的关键区别在于：FSRS 的 S 在每个物品上是统一的全局参数，'
    '而 CSM 的 d 是每条记忆独立的自适应参数（见 §2.4）。'
)

doc.add_heading('2.2 间隔强化函数', level=2)

add_body(
    '传统 SM-2 算法的强化公式为新的 EF 因子等于旧 EF 乘以固定系数，'
    '不区分复习时机。FSRS 的核心贡献是在稳定性（Stability）更新中引入 '
    '复习时的可回忆概率 R_review：'
)

add_equation('S_new = S_old × (1 + w × (1 − R_review)^α)')

add_body(
    '其中 w 为全局学习率，α ≈ 1.4 为间隔效应指数。'
    '该公式体现了间隔效应的核心洞见：复习越晚（R_review 越低），'
    '获得的长期稳定性增益越大。在 R_review = 1.0（刚学完立即复习）时增益为零，'
    '在 R_review → 0 时增益最大（但接近遗忘，实际收益-成本比在 R ≈ 0.9 处最优）。'
)

add_body(
    'CSM 将此思想适配到 Agent 场景。由于 LLM Agent 的记忆"使用"本质上是一次被动检索'
    '（记忆被检索到且被 LLM 引用），不存在主动复习调度，'
    '因此 CSM 的强化函数同时更新当前强度和长期衰减率：'
)

add_equation('R = current_strength(memory)')
add_equation('f_space(R) = (1 − R)^1.4')
add_equation('D = max(0.05, 1 − trust)')
add_equation('g = min(1.0, 0.15 × f_space(R) × D)')
add_equation('S_new = S_old × (1 + g)')
add_equation('d_new = ln(2) / S_new')
add_equation('s_new = 1.0')

add_body(
    '其中 f_space(R) 为间隔效应因子，在 R = 0.5 处取得最大值约 0.38；'
    'D 为难度因子，将 trust（信任度）映射为难度（信任度高的记忆"简单"，增益小）；'
    'g 为综合增益；S_old = ln(2)/d_old 为旧半衰期；'
    'd_new 为更新后的衰减率。强化后的当前强度 s_new 始终重置为 1.0，'
    '体现"刚复习完完全记得"的语义。'
)

add_body(
    '表 1 展示了不同初始状态下间隔效应的数值表现。'
    '当记忆在 30 天未访问后首次被使用（R ≈ 0.33），'
    '其衰减率从 0.020 降至 0.019，半衰期延长约 3.5 天。'
    '若记忆刚被使用（R ≈ 1.0），强化几乎无效（g ≈ 0），'
    '有效避免了频繁使用导致的过度强化。'
)

# 表 1
table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
headers = ['场景', '初始 R', 'f_space(R)', 'gain', '新 d']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

data = [
    ['刚使用 (0 天)', '1.000', '0.000', '0.000', '0.02000'],
    ['1 天后使用', '0.980', '≈0', '≈0', '0.02000'],
    ['10 天后使用', '0.409', '0.479', '0.036', '0.01931'],
    ['30 天后使用', '0.329', '0.588', '0.044', '0.01918'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val
        for p in table.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('表 1：不同复习间隔下的间隔效应（初始强度 0.6，trust=0.5，d=0.02）')
run.font.size = Pt(9)
run.font.italic = True

doc.add_heading('2.3 检索模型', level=2)

add_body(
    'CSM 采用混合检索策略：语义嵌入为主，FTS5 关键词为辅。'
    '对查询 q 和记忆 m，语义相似度由 BGE-large-zh-v1.5 的归一化余弦相似度给出：'
)

add_equation('sim(q, m) = max(0, Σᵢ qᵢ · mᵢ)')

add_body('最终检索得分为三项的乘积：')

add_equation('score(q, m) = sim(q, m) × R(m) × (1 + boost(m))')

add_body(
    '其中 R(m) 为记忆的当前强度（衰减后），boost(m) 为进化引擎维护的经验偏置。'
    '该公式确保（1）语义相关的记忆优先；（2）高频使用的记忆优先（通过 R 和 boost 体现）；'
    '（3）被反复纠正的记忆降权（boost 可能为负）。'
)

add_body(
    '当关键词匹配存在时（FTS5 BM25），最终得分取语义得分和关键词驱动的'
    '替代得分（0.3 × keyword_score × R × (1+boost)）的最大值，'
    '以捕获语义模型可能遗漏的精确术语匹配。'
)

doc.add_heading('2.4 自适应进化模型', level=2)

add_body(
    'CSM 维护每条记忆的三个自适应参数：衰减率（decay_rate, d）、'
    '检索偏置（boost, b）和信任度（trust, τ）。'
    '每轮对话结束后，进化引擎分析 LLM 对已检索记忆的使用情况，'
    '自动调整这些参数。'
)

add_body('（1）使用反馈（used）：记忆被 LLM 实际引用。')

add_equation('b_new = min(1.0, b + 0.05)')
add_equation('τ_new = min(0.98, τ + 0.03 × (1 − R))')

add_body('注意：信任度的更新幅度与当前 R 负相关——在 R 较低时正确"回忆"出该记忆，'
         '说明该记忆具有更高的长期可靠性，因此信任度提升更大。'
         '这与 FSRS 中 Difficulty 下调逻辑一致。')

add_body('（2）无视反馈（ignored）：记忆被检索到但未被 LLM 引用。')

add_equation('penalty = 0.02 × (0.3 + 0.7 × R)')
add_equation('b_new = max(−0.8, b − penalty)')

add_body(
    '惩罚因子与 R 正相关：当 R 高（记忆新鲜）时被无视，说明该记忆确实不相关，'
    '惩罚更大（0.02）；当 R 低（几乎遗忘）时被无视，可能只是"没想起来"，'
    '惩罚更小（0.006）。'
)

add_body('（3）纠正反馈（corrected）：用户明确指出记忆内容有误。')

add_equation('b_new = max(−0.8, b − 0.2)')
add_equation('τ_new = max(0.05, τ × 0.7)')
add_equation('d_new = min(0.3, d × 1.5)')

add_body(
    '纠正触发快速遗忘：衰减率放大 1.5 倍，信任度降至 70%。'
    '同时错误计数器 error_count 加一，供后续分析使用。'
)

doc.add_heading('2.5 动态分层与归档', level=2)

add_body(
    'CSM 使用动态百分位阈值将记忆分为四层：L1（前 20%）、L2（20%-60%）、'
    'L3（60%-90%）和 COLD（后 10%）。阈值由当前活跃记忆的强度分布动态计算：'
)

add_equation('L1_threshold = min(0.85, percentile(0.20))')
add_equation('L3_threshold = min(0.50, percentile(0.90))')

add_body(
    'L1 上限保护（0.85）确保即使极端分布下高强度记忆始终为 L1；'
    'L3 上限保护（0.50）确保新记忆（强度 0.6）初始至少为 L3。'
)

add_body(
    '睡眠整理（sleep_consolidate）采用 FSRS 的绝对 R 阈值策略：'
    '当记忆的可回忆概率 R < 0.01 时自动归档，'
    '归档后的记忆不参与回答注入检索（但仍可见于写入仲裁模式）。'
    '此举避免了传统百分位归档的缺陷——即使所有记忆都很弱时也会强制保留大部分记忆。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. 系统架构
# ════════════════════════════════════════════════════════════
doc.add_heading('3. 系统架构', level=1)

doc.add_heading('3.1 模块设计', level=2)

add_body(
    'CSM 由 14 个 Python 模块组成，核心模块职责如下：'
    '（1）engine.py：记忆生命周期管理，包括添加（含去重检测）、强化、睡眠整理和健康报告；'
    '（2）store.py：SQLite 持久化层，含 FTS5 虚拟表和 schema 迁移；'
    '（3）strength.py：指数衰减、FSRS 风格间隔强化和动态百分位分层；'
    '（4）retrieval.py：混合检索器，实现语义-FTS5 双路检索和答案注入门控；'
    '（5）evolution.py：自适应进化引擎，实现反馈检测和参数自适应；'
    '（6）embedding.py：BGE-large-zh-v1.5 本地嵌入后端；'
    '（7）extractor.py：DeepSeek LLM 仲裁器，包含 JSON schema 校验和安全策略；'
    '（8）adapters.py：pi Agent、OpenClaw 和 Hermes 的集成适配层；'
    '（9）server.py：HTTP Sidecar 服务与内置 Web 管理控制台。'
)

doc.add_heading('3.2 记忆生命周期', level=2)

add_body(
    '一条记忆从创建到归档经历以下阶段：'
    '（1）创建：LLM 仲裁器从对话中提取记忆，或用户通过 /remember 命令显式存入；'
    '初始强度 0.6，衰减率 0.02/天。'
    '（2）去重：新记忆通过语义相似度（≥0.92）和词汇重叠度（≥0.45）双重阈值与已有记忆比较，'
    '重复则合并而非新建。'
    '（3）检索：每次 Agent 对话前，CSM 检索相关记忆并注入到系统提示中；检索本身不触发强化。'
    '（4）强化：记忆在对话中被 LLM 实际引用后触发强化，调用的 reinforce() 函数'
    '根据当前可回忆概率计算间隔效应增益并更新衰减率。'
    '（5）进化：进化引擎分析反馈（used/ignored/corrected），调整 boost/trust/decay_rate。'
    '（6）归档：睡眠整理时，R < 0.01 的记忆自动归档。'
)

doc.add_heading('3.3 安全与隐私', level=2)

add_body(
    'CSM 实现三层安全保护。第一层：敏感度自动标注——检测 API Key、邮箱、手机号等模式，'
    '标记为 secret/personal 级别，但不修改原内容。'
    '第二层：作用域三态分类——记忆按内容自动归入个人（personal）、项目共享（project）、'
    '或全局（global）作用域，个人记忆仅对当前用户可见，全局记忆跨项目去重和检索。'
    '第三层：HTTP API Key 认证——Sidecar 服务支持 X-CSM-API-Key 和 Bearer Token 两种认证方式。'
)

doc.add_heading('3.4 部署与集成', level=2)

add_body(
    'CSM 以零外部依赖的 pip 包形式发布（可选依赖 sentence-transformers 提供语义搜索）。'
    '安装后通过 membrain serve 命令启动 HTTP Sidecar 和 Web 管理控制台。'
    'pi Agent 通过 pi-extension/mb-memory.ts 自动集成：'
    '扩展启动时自动拉起 Sidecar 子进程，每次对话前检索记忆注入上下文，'
    '每次对话后通过 DeepSeek LLM 仲裁器提取记忆。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. 实验评估
# ════════════════════════════════════════════════════════════
doc.add_heading('4. 实验评估', level=1)

doc.add_heading('4.1 强度模型评测', level=2)

add_body(
    '为验证强度模型的数学正确性，构建了 16 个测试用例，覆盖衰减（7 个）、'
    '强化（4 个）、综合（2 个）和动态阈值（3 个）四个维度。'
    '所有衰减测试均通过（通过率 100%），验证了指数衰减公式在 1 天至 365 天'
    '各时间尺度下的正确性。强化测试中，间隔效应因子使 reinforce() 始终返回 1.0'
    '（刚强化完的可回忆概率），而衰减率的更新量随复习间隔增大而增大，'
    '验证了间隔效应的预期行为。动态阈值测试验证了百分位算法在标准分布、'
    '极小样本（n<5 回退固定阈值）和极端分布下的正确性。'
)

doc.add_heading('4.2 检索评测', level=2)

add_body(
    '检索评测使用 18+ 个检索测试用例，指标包括 Recall@k、Precision@k、'
    'MRR（Mean Reciprocal Rank）和 NDCG（Normalized Discounted Cumulative Gain）。'
    '测试数据覆盖项目依赖变更（Supersede 链）、用户偏好、'
    '临时信息和跨项目隔离等场景。'
    'Recall@k ≥ 0.6，Forbidden Hit Rate ≤ 0.2，验证了混合检索策略的有效性和'
    '作用域隔离的正确性。'
)

doc.add_heading('4.3 端到端评测', level=2)

add_body(
    '端到端评测使用 10+ 个历史对话序列，模拟多轮 Agent 交互后验证记忆系统是否'
    '正确提取和检索了持久化信息。准确率 ≥ 0.8，记忆污染率 ≤ 0.3。'
    '评测包括偏好持久化、纠正检测、临时信息过滤和冲突写入消解等场景。'
)

doc.add_heading('4.4 嵌入质量评测', level=2)

add_body(
    '语义嵌入质量评测测试了 BGE-large-zh-v1.5 在同义词召回、改写召回和跨语言召回'
    '三个维度的表现。同义词测试包括"安装依赖"→ "bun install" 等 3 组配对；'
    '改写测试包括"回答有什么要求"→ "简洁" 等 3 组配对；'
    '跨语言测试包括"如何部署服务"→ "Deploy with Docker Compose" 等 3 组配对。'
    '实验表明 BGE-large-zh-v1.5 在中文语义匹配上表现优异，平均语义相似度达到较高水平。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. 相关工作
# ════════════════════════════════════════════════════════════
doc.add_heading('5. 相关工作', level=1)

add_body(
    '间隔重复系统。Ebbinghaus (1885) 首次定量描述了人类记忆的遗忘曲线。'
    'SuperMemo 的 SM-2 算法 [Wozniak, 1990] 将 Easiness Factor 引入间隔重复，'
    '成为 Anki 等闪卡软件的基础。FSRS [Ye, 2022] 从 SM-2 演进为三状态 DSR 模型'
    '（Difficulty, Stability, Retrievability），通过随机最短路径优化学习参数。'
    'CSM 与 FSRS 的关键区别在于：（1）CSM 面向 LLM Agent 的隐式反馈而非人类显式评分；'
    '（2）CSM 的"复习"是被动检索触发而非主动调度。'
)

add_body(
    'LLM 记忆系统。MemGPT [Packer et al., 2023] 通过操作系统的虚拟内存抽象'
    '管理 LLM 上下文，但依赖全量上下文而非选择性的长期存储。'
    'Mem0 提供了基于向量的记忆检索，但缺乏强度衰减和自适应进化机制。'
    'LangChain Memory 提供了基础的对话历史存储，但不支持语义去重和自动仲裁。'
    'CSM 的独特贡献在于将间隔重复的数学严谨性引入 LLM 记忆管理。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. 结论与展望
# ════════════════════════════════════════════════════════════
doc.add_heading('6. 结论与展望', level=1)

add_body(
    '本文提出了连续强度记忆系统（CSM），一种为 LLM Agent 设计的长期记忆子系统。'
    'CSM 的核心理念是"从简单算法中涌现智能行为"——整个系统仅依赖指数衰减和间隔强化'
    '两条规则，通过与 LLM 的交互反馈自适应进化，实现了语义检索、'
    '自动仲裁、冲突检测和隐私隔离等高级功能。'
)

add_body(
    '本文的主要贡献包括：（1）将 FSRS 间隔效应理论适配到 Agent 场景，'
    '提出了基于隐式反馈的自适应记忆模型；（2）设计了时间感知的强化函数，'
    '使记忆的长期保留取决于使用间隔而非单纯的使用次数；'
    '（3）构建了包含 auto-merge、LLM 仲裁、进化反馈的完整记忆生命周期；'
    '（4）以零外部依赖的 pip 包形式开源发布。'
)

add_body(
    '未来工作方向包括：（1）引入向量索引（如 sqlite-vec）以支持大规模记忆检索；'
    '（2）基于用户数据训练个性化的 FSRS 参数（w₈, w₉, w₁₀）；'
    '（3）扩展 LLM 仲裁器支持更多模型提供商；'
    '（4）增加记忆图谱功能（基于 memory_links 的关系网络探索）。'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 参考文献
# ════════════════════════════════════════════════════════════
doc.add_heading('参考文献', level=1)

refs = [
    '[1] Ebbinghaus, H. (1885). Memory: A Contribution to Experimental Psychology.',
    '[2] Wozniak, P. A. (1990). Optimization of Learning. Master\'s Thesis, Poznan University of Technology.',
    '[3] Ye, J. et al. (2022). FSRS: A Free Spaced Repetition Scheduler. https://github.com/open-spaced-repetition/fsrs4anki',
    '[4] Packer, C. et al. (2023). MemGPT: Towards LLMs as Operating Systems. arXiv:2310.08560.',
    '[5] Liu, N. F. et al. (2023). Lost in the Middle: How Language Models Use Long Contexts. arXiv:2307.03172.',
    '[6] Wozniak, P. A. & Gorzelanczyk, E. J. (1994). Optimization of Repetition Spacing in the Practice of Learning. Acta Neurobiologiae Experimentalis, 54:59-62.',
    '[7] Reimers, N. & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. arXiv:1908.10084.',
    '[8] Xiao, S. et al. (2023). BGE: C-Pack: Packaged Resources To Advance General Chinese Embedding. arXiv:2309.07597.',
    '[9] SQLite Consortium. (2024). SQLite FTS5 Extension. https://www.sqlite.org/fts5.html',
    '[10] DeepSeek-AI. (2025). DeepSeek-V4 Technical Report. https://api.deepseek.com',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(10)

# ── 保存 ──────────────────────────────────────────────────
output_path = os.path.expanduser('~/Desktop/CSM_学术论文.docx')
doc.save(output_path)
print(f'论文已生成: {output_path}')
